#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""按主题文件夹归档对话：从 stdin 读 Markdown，生成 .md 与 .docx 两个版本。

用法:
  python3 archive.py --root "根目录" --folder "主题文件夹" --name "提问前10字"
  （Markdown 内容从标准输入读取；输出到 <root>/<folder>/<name>.md 和 <name>.docx）
"""
import os, sys, re, argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EA_FONT = "PingFang SC"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x59, 0x59, 0x59)

def set_run_font(run, name=EA_FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn('w:eastAsia'), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None: run.font.color.rgb = color

def shade_paragraph(paragraph, fill="F5F5F5"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def shade_cell(cell, fill="DEEAF6"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_inline(paragraph, text):
    for part in re.split(r'(`[^`]+`)', text):
        if not part: continue
        if part.startswith('`') and part.endswith('`') and len(part) >= 2:
            r = paragraph.add_run(part[1:-1]); set_run_font(r, MONO_FONT, size=10.5, color=RGBColor(0xC0,0x25,0x32))
            continue
        pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]*\))')
        for tok in pattern.split(part):
            if not tok: continue
            if tok.startswith('**') and tok.endswith('**') and len(tok) > 4:
                r = paragraph.add_run(tok[2:-2]); set_run_font(r, bold=True)
            elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
                r = paragraph.add_run(tok[1:-1]); set_run_font(r, italic=True)
            elif tok.startswith('[') and '](' in tok:
                m = re.match(r'\[([^\]]+)\]\([^)]*\)', tok)
                r = paragraph.add_run(m.group(1)); set_run_font(r, color=ACCENT); r.underline = True
            else:
                r = paragraph.add_run(tok); set_run_font(r)

def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    sizes = {1: 17, 2: 14, 3: 12.5, 4: 11.5}
    set_run_font(r, EA_FONT, size=sizes.get(level, 11), bold=True, color=ACCENT if level <= 3 else RGBColor(0x33,0x33,0x33))
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(6)

def add_paragraph(doc, text, quote=False):
    p = doc.add_paragraph()
    if quote:
        p.paragraph_format.left_indent = Inches(0.25)
        add_inline(p, text)
        for r in p.runs: r.font.color.rgb = GRAY; r.italic = True
    else:
        add_inline(p, text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

def add_list_item(doc, text, ordered=False, task=None):
    p = doc.add_paragraph(style='List Number' if ordered else 'List Bullet')
    if task is not None:
        r = p.add_run('☑ ' if task else '☐ '); set_run_font(r, EA_FONT, size=11)
    add_inline(p, text)
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.15

def add_code_block(doc, code):
    for line in code.split('\n'):
        p = doc.add_paragraph()
        r = p.add_run(line if line else ' ')
        set_run_font(r, MONO_FONT, size=9.5, color=RGBColor(0x33,0x33,0x33))
        shade_paragraph(p)
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

def _w(s): return sum(2 if ord(c) > 0x2E80 else 1 for c in s)

def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols); table.style = 'Table Grid'; table.autofit = False
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j); cell.text = ''
            txt = row[j] if j < len(row) else ''
            p = cell.paragraphs[0]; add_inline(p, txt.strip())
            for r in p.runs: set_run_font(r, EA_FONT, size=10, bold=(i == 0))
            if i == 0: shade_cell(cell)
    usable = 6.47
    weights = [max(max(_w(rows[i][j]) if j < len(rows[i]) else 0 for i in range(len(rows))), 6) for j in range(ncols)]
    widths = [max(0.7, usable * w / sum(weights)) for w in weights]
    scale = usable / sum(widths); widths = [w * scale for w in widths]
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(sum(widths) * 1440))); tblW.set(qn('w:type'), 'dxa')
    for row in table.rows:
        for j, cell in enumerate(row.cells): cell.width = Inches(widths[j])
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:space'), '1'); b.set(qn('w:color'), 'BFBFBF')
    pbdr.append(b); pPr.append(pbdr); p.paragraph_format.space_after = Pt(8)

def convert_markdown(doc, md_text):
    lines = md_text.split('\n'); i = 0
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith('```'):
            i += 1; buf = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1; add_code_block(doc, '\n'.join(buf)); continue
        if s.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i+1]):
            header = [c.strip() for c in s.strip('|').split('|')]; i += 2; body = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                body.append([c.strip() for c in lines[i].strip().strip('|').split('|')]); i += 1
            add_table(doc, [header] + body); continue
        if re.match(r'^\s*(\*\*\*|---|___)\s*$', s): add_hr(doc); i += 1; continue
        m = re.match(r'^(#{1,6})\s+(.*)$', s)
        if m: add_heading(doc, m.group(2).strip(), max(1, min(6, len(m.group(1))))); i += 1; continue
        if s.startswith('>'):
            add_paragraph(doc, s.lstrip('> ').strip(), quote=True); i += 1; continue
        mt = re.match(r'^[-*]\s+\[([ xX])\]\s+(.*)$', s)
        if mt: add_list_item(doc, mt.group(2), task=(mt.group(1).lower() == 'x')); i += 1; continue
        if re.match(r'^[-*]\s+', s): add_list_item(doc, re.sub(r'^[-*]\s+', '', s)); i += 1; continue
        mo = re.match(r'^\d+[.)]\s+(.*)$', s)
        if mo: add_list_item(doc, mo.group(1), ordered=True); i += 1; continue
        if not s: i += 1; continue
        add_paragraph(doc, s); i += 1

def build_docx(md_text):
    doc = Document()
    normal = doc.styles['Normal']; normal.font.name = EA_FONT; normal.font.size = Pt(11)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), EA_FONT)
    for sec in doc.sections:
        sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
        sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
    convert_markdown(doc, md_text)
    return doc

def safe_component(s):
    s = re.sub(r'[\\/:*?"<>|\s]+', '', s)   # 去掉非法字符和空格
    s = s.strip('.')
    return s or '未命名'

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--root', required=True)
    ap.add_argument('--folder', required=True)
    ap.add_argument('--name', required=True)
    args = ap.parse_args()
    md_text = sys.stdin.read()
    folder = safe_component(args.folder)
    name = safe_component(args.name)
    out_dir = os.path.join(args.root, folder)
    os.makedirs(out_dir, exist_ok=True)
    md_path = os.path.join(out_dir, f'{name}.md')
    docx_path = os.path.join(out_dir, f'{name}.docx')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_text)
    build_docx(md_text).save(docx_path)
    print('SAVED_MD=' + md_path)
    print('SAVED_DOCX=' + docx_path)

if __name__ == '__main__':
    main()
